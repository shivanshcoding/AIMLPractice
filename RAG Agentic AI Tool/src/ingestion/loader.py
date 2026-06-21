"""
Multi-Format Document Loader.

Loads documents from various file formats: PDF, DOCX, TXT, Markdown, HTML.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

import structlog

from src.core.models import Document, DocumentMetadata
from src.core.exceptions import IngestionError

logger = structlog.get_logger(__name__)


class DocumentLoader:
    """Load documents from files into the Document model."""

    _SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".html", ".htm", ".csv"}

    def load_file(self, file_path: str | Path, metadata: dict[str, Any] | None = None) -> Document:
        """Load a single file into a Document."""
        path = Path(file_path)
        if not path.exists():
            raise IngestionError(f"File not found: {path}")

        ext = path.suffix.lower()
        if ext not in self._SUPPORTED_EXTENSIONS:
            raise IngestionError(f"Unsupported file type: {ext}")

        content = self._read_file(path)

        doc_metadata = DocumentMetadata(
            source=str(path),
            file_type=ext,
            file_size_bytes=path.stat().st_size,
            **(metadata or {}),
        )

        return Document(content=content, metadata=doc_metadata)

    def load_directory(
        self, dir_path: str | Path, metadata: dict[str, Any] | None = None
    ) -> list[Document]:
        """Load all supported files from a directory."""
        directory = Path(dir_path)
        if not directory.is_dir():
            raise IngestionError(f"Not a directory: {directory}")

        documents = []
        for file_path in sorted(directory.rglob("*")):
            if file_path.suffix.lower() in self._SUPPORTED_EXTENSIONS:
                try:
                    doc = self.load_file(file_path, metadata)
                    documents.append(doc)
                except Exception as e:
                    logger.warning("file_load_failed", path=str(file_path), error=str(e))

        logger.info("directory_loaded", path=str(directory), documents=len(documents))
        return documents

    def load_text(self, text: str, metadata: dict[str, Any] | None = None) -> Document:
        """Load raw text as a Document."""
        doc_metadata = DocumentMetadata(**(metadata or {}))
        return Document(content=text, metadata=doc_metadata)

    def _read_file(self, path: Path) -> str:
        """Read file content based on extension."""
        ext = path.suffix.lower()

        if ext == ".pdf":
            return self._read_pdf(path)
        elif ext == ".docx":
            return self._read_docx(path)
        elif ext in (".txt", ".md", ".csv"):
            return path.read_text(encoding="utf-8")
        elif ext in (".html", ".htm"):
            return self._read_html(path)
        else:
            return path.read_text(encoding="utf-8")

    def _read_pdf(self, path: Path) -> str:
        """Read PDF using pypdf."""
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(path))
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n\n".join(pages)
        except ImportError:
            raise IngestionError("pypdf not installed. Run: pip install pypdf")

    def _read_docx(self, path: Path) -> str:
        """Read DOCX using python-docx."""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(str(path))
            paragraphs = [p.text for p in doc.paragraphs]
            return "\n\n".join(paragraphs)
        except ImportError:
            raise IngestionError("python-docx not installed. Run: pip install python-docx")

    def _read_html(self, path: Path) -> str:
        """Read HTML and extract text."""
        try:
            from bs4 import BeautifulSoup
            html = path.read_text(encoding="utf-8")
            soup = BeautifulSoup(html, "html.parser")
            return soup.get_text(separator="\n\n", strip=True)
        except ImportError:
            raise IngestionError("beautifulsoup4 not installed.")
